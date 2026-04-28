"""Extract .docx content to markdown for use as skill reference files."""
import sys
from pathlib import Path
from docx import Document

def extract_docx(input_path: str, output_path: str) -> None:
    doc = Document(input_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = para.style.name
        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        else:
            lines.append(text)
    for table in doc.tables:
        header = "| " + " | ".join(cell.text.strip() for cell in table.rows[0].cells) + " |"
        separator = "| " + " | ".join("---" for _ in table.rows[0].cells) + " |"
        rows = []
        for row in table.rows[1:]:
            rows.append("| " + " | ".join(cell.text.strip() for cell in row.cells) + " |")
        lines.extend(["", header, separator] + rows + [""])
    Path(output_path).write_text("\n".join(lines), encoding="utf-8")
    print(f"Extracted: {output_path}")

if __name__ == "__main__":
    extract_docx(sys.argv[1], sys.argv[2])
