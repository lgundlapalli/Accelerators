"""
Extract text from .docx files using python-docx.
Falls back to macOS textutil for .doc files.

Usage: python3 extract_doc.py <file_path>
"""

import sys
import os
import subprocess


def extract_docx(path):
    """Extract text from .docx using python-docx."""
    from docx import Document
    doc = Document(path)

    sections = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Preserve heading hierarchy
        if para.style and para.style.name.startswith('Heading'):
            level = para.style.name.replace('Heading ', '').strip()
            try:
                level = int(level)
                sections.append(f"\n{'#' * level} {text}\n")
            except ValueError:
                sections.append(f"\n## {text}\n")
        else:
            sections.append(text)

    # Also extract tables
    for i, table in enumerate(doc.tables):
        sections.append(f"\n[Table {i + 1}]")
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            sections.append(" | ".join(cells))

    return "\n".join(sections)


def extract_doc_textutil(path):
    """Extract text from .doc using macOS textutil."""
    result = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", path],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print(f"Error: textutil failed — {result.stderr}", file=sys.stderr)
        sys.exit(1)
    return result.stdout


def main():
    if len(sys.argv) != 2:
        print("Usage: python3 extract_doc.py <file_path>", file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.exists(path):
        print(f"Error: file not found — {path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(path)[1].lower()

    if ext == '.docx':
        print(extract_docx(path))
    elif ext == '.doc':
        print(extract_doc_textutil(path))
    else:
        print(f"Error: unsupported format '{ext}'. Use .doc or .docx", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
