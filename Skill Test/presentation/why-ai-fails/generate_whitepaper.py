"""
White Paper: The Cultural Debt of AI Transformation
====================================================
Generates a formatted .docx white paper using python-docx.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ============================================================
# STYLES & FORMATTING
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Calibri'
style_h1.font.size = Pt(22)
style_h1.font.color.rgb = RGBColor(0x00, 0x00, 0x50)
style_h1.font.bold = True

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Calibri'
style_h2.font.size = Pt(16)
style_h2.font.color.rgb = RGBColor(0x00, 0x00, 0x50)
style_h2.font.bold = True

style_h3 = doc.styles['Heading 3']
style_h3.font.name = 'Calibri'
style_h3.font.size = Pt(13)
style_h3.font.color.rgb = RGBColor(0x33, 0x33, 0xAA)
style_h3.font.bold = True


def add_para(text, bold=False, italic=False, size=11, space_after=6, alignment=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x50)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        # Clear default and add run
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para("WHITE PAPER", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
         color=RGBColor(0x33, 0x33, 0xAA), space_after=12)
add_para("The Cultural Debt of AI Transformation", bold=True, size=28,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x00, 0x00, 0x50), space_after=8)
add_para("Why AI Initiatives Fail — and Why the Fix Isn't Technology",
         italic=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
         color=RGBColor(0x33, 0x33, 0xAA), space_after=24)

add_para("April 2026", size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Proprietary and confidential — do not distribute",
         italic=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
         color=RGBColor(0x88, 0x88, 0x88))

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Contents', level=1)
toc_items = [
    "1. Abstract",
    "2. Problem Statement",
    "3. Why Cultural Debt Matters",
    "4. Why Now",
    "5. Key Considerations",
    "6. Remediations",
    "7. Path Forward",
    "8. References",
]
for item in toc_items:
    add_para(item, size=12, space_after=4)

doc.add_page_break()

# ============================================================
# 1. ABSTRACT
# ============================================================
doc.add_heading('1. Abstract', level=1)

add_para(
    "Artificial intelligence is the defining technology investment of the decade. Global AI spending "
    "is projected to reach $4.6 trillion by 2027. Yet 70–85% of AI initiatives fail to deliver their "
    "expected value — and technology is rarely the cause. Across industries, geographies, and company "
    "sizes, the primary driver of AI failure is organizational and cultural resistance: the accumulated "
    "mistrust, misaligned incentives, skill gaps, and eroded psychological safety that grow with every "
    "unsuccessful attempt."
)
add_para(
    "This paper introduces the concept of cultural debt — the organizational parallel to technical "
    "debt. Where technical debt represents shortcuts in code that create future rework, cultural debt "
    "represents shortcuts in change management that create future resistance. Unlike technical debt, "
    "cultural debt is untracked, unmeasured, and unbudgeted. It compounds quarterly. And it is the "
    "single largest predictor of whether an AI initiative will succeed or fail."
)
add_para(
    "Drawing on research from McKinsey, Gartner, MIT CISR, Harvard Business Review, and the Six "
    "Sources of Influence behavioral science framework, this paper examines why cultural debt "
    "accumulates, why the cost of inaction is rising, and what organizations can do — starting this "
    "quarter — to begin paying it down. The recommendations are structured around three actionable "
    "storylines: Culture, Change Management, and Leadership."
)

# ============================================================
# 2. PROBLEM STATEMENT
# ============================================================
doc.add_heading('2. Problem Statement', level=1)

add_para(
    "The AI industry has a failure rate that would be unacceptable in any other domain of capital "
    "investment. The numbers are consistent and well-documented:"
)

# Key stats table
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
headers = ["Metric", "Value", "Source"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

data = [
    ["AI initiatives failing to deliver expected value", "70–85%", "McKinsey (2023); Gartner (2024)"],
    ["Failures attributed to technology limitations", "< 15%", "BCG AI Adoption Survey (2024)"],
    ["#1 cited reason for failure", "Organizational and cultural resistance", "Rand Corporation (2024)"],
    ["Companies reporting significant financial benefit from AI", "10%", "BCG & MIT Sloan (2024)"],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

add_para("", space_after=8)

add_para(
    "The pattern is clear: the technology works. Large language models generate accurate output. "
    "Machine learning platforms scale. Cloud infrastructure is mature. What fails is the human "
    "system surrounding the technology — the organization's ability to adopt, integrate, and sustain "
    "AI-driven ways of working."
)

add_para(
    "This paper argues that the root cause is not a lack of change management, executive sponsorship, "
    "or training programs — most failing organizations have all three. The root cause is cultural "
    "debt: the invisible, compounding residue of every prior technology initiative that was deployed "
    "without adequately investing in the humans expected to adopt it."
)

# ============================================================
# 3. WHY CULTURAL DEBT MATTERS
# ============================================================
doc.add_heading('3. Why Cultural Debt Matters', level=1)

doc.add_heading('3.1 Defining Cultural Debt', level=2)

add_para(
    "Technical debt is a concept every engineering leader understands: shortcuts taken in code that "
    "create future rework. It is visible in code reviews, measurable in build times and bug rates, "
    "and routinely budgeted for in sprint planning. Cultural debt is the organizational equivalent."
)

# Comparison table
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = "Technical Debt"
table.rows[0].cells[1].text = "Cultural Debt"
for p in table.rows[0].cells[0].paragraphs:
    p.runs[0].bold = True
for p in table.rows[0].cells[1].paragraphs:
    p.runs[0].bold = True

comparisons = [
    ("Shortcuts in code that create future rework", "Shortcuts in change management that create future resistance"),
    ("Visible in code reviews, build times, bug rates", "Invisible until the next initiative stalls"),
    ("Tracked, measured, budgeted for", "Untracked, unmeasured, unbudgeted"),
    ("Compounds as system complexity grows", "Compounds as organizational cynicism grows"),
    ("Fix: refactor the code", "Fix: rebuild trust, realign incentives, invest in people"),
]
for r, (tech, cult) in enumerate(comparisons):
    table.rows[r + 1].cells[0].text = tech
    table.rows[r + 1].cells[1].text = cult

add_para("", space_after=8)

add_quote(
    "Cultural debt accrues every time an organization deploys technology without investing in the "
    "humans who must adopt it."
)

doc.add_heading('3.2 The Compounding Cycle', level=2)

add_para(
    "Cultural debt does not sit idle. It compounds with every failed initiative, following a "
    "predictable four-stage cycle:"
)

stages = [
    ("Cycle 1 — Healthy Skepticism: ", '"Let\'s see if this works." Resistance is low and rational. '
     "Teams are willing to experiment but cautious. The cost to address cultural readiness at this "
     "stage is a leadership conversation and a readiness assessment."),
    ("Cycle 2 — Active Resistance: ", '"We tried this before." Previous failures are cited as evidence '
     "that AI doesn't work here. Middle managers adopt a wait-and-see posture. Shadow processes "
     "emerge as teams route around the new tools."),
    ("Cycle 3 — Organizational Antibodies: ", '"AI isn\'t for us." Resistance becomes structural and '
     "unconscious. Shelly Palmer describes these as 'bureaucratic antibodies' — institutional immune "
     "responses that narrow scope, expand approval criteria, and declare initiatives 'premature.' "
     "The cost to address cultural debt at this stage requires a visible organizational reset."),
    ("Cycle 4 — Structural Lock-in: ", "The best people have stopped engaging. Workarounds are "
     "institutionalized. The narrative that 'AI doesn't deliver' is accepted as fact. Addressing "
     "cultural debt at this stage requires organizational redesign — not a workshop."),
]
for prefix, text in stages:
    add_bullet(text, bold_prefix=prefix)

add_para(
    "Each cycle is not merely a repeat — it is worse. The cynicism is deeper, the shadow processes "
    "are more entrenched, and the talent most capable of driving adoption has disengaged. Research "
    "shows that second AI attempts face approximately 3x higher resistance when cultural debt from "
    "the first attempt goes unaddressed."
)

doc.add_heading('3.3 Five Cultural Patterns That Predict AI Failure', level=2)

add_para(
    "Five recurring organizational patterns serve as diagnostic indicators of cultural debt. If "
    "two or more are present, the initiative is at risk before a single model is trained:"
)

patterns = [
    ("Permission Culture: ", "Teams wait for explicit approval before experimenting with AI. "
     "Innovation dies in committee while competitors move faster. Experimentation requires a "
     "business case before anyone has learned enough to write one."),
    ("Metric Misalignment: ", "AI success is measured by deployment — models in production, "
     "platforms launched — rather than adoption or business impact. Teams ship solutions no one "
     "uses and declare victory based on the wrong scorecard."),
    ("The Invisible Middle: ", "Middle management neither actively sponsors nor visibly blocks "
     "AI initiatives — they wait. This is the least discussed and often most lethal pattern. "
     "Middle managers control execution. Their rational response to ambiguous leadership signals "
     "is inaction."),
    ("Skills Theater: ", "Training programs check compliance boxes but do not build real "
     "capability. Teams attend AI workshops on Friday and return to pre-AI workflows on Monday. "
     "The organization reports '500 employees trained in AI' while actual usage remains flat."),
    ("Trust Deficit: ", "The narrative that 'AI will replace us' circulates unchallenged. "
     "Leadership has not articulated a credible vision of Human + AI collaboration. The best "
     "talent disengages or leaves. Adoption becomes coercion rather than conviction."),
]
for prefix, text in patterns:
    add_bullet(text, bold_prefix=prefix)

# ============================================================
# 4. WHY NOW
# ============================================================
doc.add_heading('4. Why Now', level=1)

doc.add_heading('4.1 The Quarterly Compounding Effect', level=2)

add_para(
    "Cultural debt is not a static problem. It compounds on a quarterly cadence, and the cost "
    "to address it roughly doubles with each stage:"
)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
headers = ["Timeline", "Stage", "Cost to Address"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

timeline_data = [
    ["Quarter 1", "Skepticism — \"Let's wait and see\"", "A leadership conversation"],
    ["Quarter 2", "Resistance — \"This didn't work last time\"", "A facilitated readiness assessment"],
    ["Quarter 4", "Antibodies — \"AI isn't for us\"", "A visible organizational reset"],
    ["Year 2", "Structural lock-in — best people stopped engaging", "A transformation program"],
]
for r, row_data in enumerate(timeline_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

add_para("", space_after=8)

add_quote(
    "Addressing cultural debt in Q1 costs a leadership conversation. "
    "Addressing it in Year 2 costs a transformation program."
)

doc.add_heading('4.2 The Investment Gap Is Widening', level=2)

add_para(
    "Global AI investment is projected to reach $4.6 trillion by 2027 (IDC). The vast majority "
    "of this capital is allocated to technology: platforms, models, infrastructure, and tooling. "
    "The fraction allocated to organizational readiness — the cultural infrastructure required "
    "to make that technology deliver value — remains negligible."
)
add_para(
    "This gap is widening, not closing. As AI capabilities advance rapidly (foundation models, "
    "autonomous agents, multimodal reasoning), the human system's ability to absorb and integrate "
    "these capabilities is falling further behind. The technology is accelerating. The "
    "organization is not."
)

doc.add_heading('4.3 The Competitive Window', level=2)

add_para(
    "Organizations that address cultural debt now gain a compounding advantage. Those that delay "
    "face a compounding deficit. The research is directional but consistent:"
)

add_bullet(
    "Organizations that invest in cultural readiness before technology deployment see "
    "2–3x higher adoption rates (MIT CISR, 2024).",
)
add_bullet(
    "Leaders who engage all six sources of influence (not just training and communication) are "
    "10x more likely to drive rapid behavioral change (MIT Sloan Management Review / VitalSmarts).",
)
add_bullet(
    "Organizations citing 'cultural readiness' as their top AI barrier increased from "
    "32% to 47% year-over-year (Deloitte State of AI, 2025). The problem is getting worse, "
    "not better.",
)

add_para(
    "The competitive window for building cultural infrastructure is now. Organizations that "
    "wait until cultural debt reaches Cycle 3 or 4 will find themselves unable to adopt AI "
    "at the speed their competitors can — regardless of how much they spend on technology."
)

# ============================================================
# 5. KEY CONSIDERATIONS
# ============================================================
doc.add_heading('5. Key Considerations', level=1)

doc.add_heading('5.1 The Six Sources of Influence Framework', level=2)

add_para(
    "The Six Sources of Influence, developed by researchers at VitalSmarts (now Crucial Learning) "
    "and validated in a 2008 MIT Sloan Management Review study, provides a diagnostic framework "
    "for understanding where behavioral adoption breaks down. The framework identifies six forces "
    "that shape human behavior, organized in a 2x3 matrix:"
)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = ""
table.rows[0].cells[1].text = "Motivation"
table.rows[0].cells[2].text = "Ability"
for p in table.rows[0].cells[1].paragraphs:
    p.runs[0].bold = True
for p in table.rows[0].cells[2].paragraphs:
    p.runs[0].bold = True

grid = [
    ["Personal", "Source 1: Do people WANT to use AI?\n(Fear, identity, trust)",
     "Source 2: Do people KNOW HOW?\n(Skills, prompting, validation)"],
    ["Social", "Source 3: Do peers and leaders MODEL it?\n(Visible role models, social proof)",
     "Source 4: Do teams SUPPORT each other?\n(Peer networks, shared practices)"],
    ["Structural", "Source 5: Do INCENTIVES reward it?\n(Promotions, KPIs, recognition)",
     "Source 6: Does the ENVIRONMENT enable it?\n(Workflow integration, tooling)"],
]
for r, (label, mot, abil) in enumerate(grid):
    row = table.rows[r + 1]
    row.cells[0].text = label
    for p in row.cells[0].paragraphs:
        p.runs[0].bold = True
    row.cells[1].text = mot
    row.cells[2].text = abil

add_para("", space_after=8)

add_para(
    "The critical insight: most AI change programs address only 2 of 6 sources — training "
    "(Source 2: Personal Ability) and communication (Source 3: Social Motivation). This leaves "
    "four sources actively pulling against adoption. Traditional change management has a 70% "
    "failure rate precisely because it operates on one-third of the forces that shape behavior."
)

doc.add_heading('5.2 Where Cultural Debt Hides in Each Source', level=2)

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Grid Accent 1'
headers = ["Source", "The Cultural Debt", "Observable Signal"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

debt_data = [
    ["1. Personal Motivation", "Fear and distrust — \"AI will replace my role\"",
     "Low voluntary usage despite access"],
    ["2. Personal Ability", "Skill gaps disguised as skepticism — \"AI isn't reliable\"",
     "Inconsistent output quality across teams"],
    ["3. Social Motivation", "No visible role models — no leader publicly using AI",
     "\"Promising but premature\" verdicts from management"],
    ["4. Social Ability", "No peer support — everyone figures it out alone",
     "Isolated pockets of adoption that don't scale"],
    ["5. Structural Motivation", "Incentives reward the old way of working",
     "Rational non-adoption by process owners"],
    ["6. Structural Ability", "AI bolted on outside core workflows",
     "High license cost, low actual usage"],
]
for r, row_data in enumerate(debt_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

add_para("", space_after=8)

add_quote(
    "If even one source is misaligned, adoption stalls. If several are misaligned, it fails entirely."
)

doc.add_heading('5.3 The People-Process-Technology Inversion', level=2)

add_para(
    "Most organizations invest in AI in the order Technology → Process → People. The evidence "
    "suggests the successful order is reversed: People → Process → Technology."
)

add_bullet(
    "People first: Redefine roles around Human + AI collaboration. Build capability through "
    "apprenticeship, not slide decks. Make experimentation safe. Address the identity threat "
    "before introducing the tool.",
    bold_prefix="People first: "
)
add_bullet(
    "Process second: Embed adoption into the delivery methodology — not as a post-launch phase. "
    "Co-design workflows with the people who will use them. Measure outcomes, not outputs.",
    bold_prefix="Process second: "
)
add_bullet(
    "Technology third: Choose platforms that meet people where they are. Integrate AI into "
    "existing workflows rather than requiring new ones. Reduce friction. The best AI tool is "
    "the one embedded in the system people already use.",
    bold_prefix="Technology third: "
)

add_para(
    "This inversion is the source of cultural debt: every organization that starts with technology "
    "selection before investing in people readiness adds to the accumulating resistance that "
    "will confront the next initiative."
)

# ============================================================
# 6. REMEDIATIONS
# ============================================================
doc.add_heading('6. Remediations', level=1)

add_para(
    "The six sources of influence map naturally to three actionable storylines. Each storyline "
    "addresses two sources and targets a specific layer of cultural debt."
)

doc.add_heading('6.1 Storyline 1: Culture (Sources 1 & 2 — Personal)', level=2)

add_para("Objective: Shift from permission culture to experimentation culture.", bold=True)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
headers = ["From (Cultural Debt)", "To (Cultural Infrastructure)", "Evidence"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

culture_data = [
    ["AI as threat — \"this will replace me\"",
     "AI as multiplier — \"this makes me better\"",
     "Energy firm: AI anomaly detection reframed from 'tool' to 'strategic safeguard' — narrative shifted adoption"],
    ["Skills theater — attend training, return to old habits",
     "Apprenticeship — learn by doing, build real capability",
     "Services firm: hands-on workshops with real scenarios drove confidence and adoption up in one quarter"],
    ["Permission culture — wait for approval to experiment",
     "Experimentation culture — safe to try, safe to fail",
     "Google Project Aristotle: psychological safety is the #1 predictor of team effectiveness"],
    ["Fear of failure — mistakes are career risk",
     "Learning from failure — mistakes are data",
     "Normalize \"I tried X and it didn't work because Y\" in team standups and retrospectives"],
]
for r, row_data in enumerate(culture_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

add_para("", space_after=4)
add_quote("Culture isn't a soft topic. It's the difference between teams who experiment and teams who hide.")

doc.add_heading('6.2 Storyline 2: Change Management (Sources 3 & 4 — Social)', level=2)

add_para("Objective: Upgrade from 2-of-6-source programs to full Six Sources adoption strategy.", bold=True)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
headers = ["From (Cultural Debt)", "To (Cultural Infrastructure)", "Evidence"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

cm_data = [
    ["Traditional CM — covers 2 of 6 sources",
     "Six Sources approach — all 6 aligned",
     "Traditional CM has a 70% failure rate. The missing 4 sources are where adoption stalls."],
    ["No visible role models — adoption feels optional",
     "Respected leaders publicly model AI use",
     "Retail org: regional leaders using AI forecasts shifted adoption from optional to standard practice"],
    ["No peer support — everyone figures it out alone",
     "Cross-functional AI councils and communities of practice",
     "Manufacturing org: joint compliance + business AI reviews reduced fear and accelerated trust"],
    ["Measure deployment, not adoption",
     "Measure adoption and impact — time-to-value, not time-to-production",
     "Microsoft: shifted KPIs from deployment to business outcome — adoption jumped 40%"],
]
for r, row_data in enumerate(cm_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

add_para("", space_after=4)
add_quote(
    "Change management isn't failing. It's incomplete. "
    "Covering 2 of 6 sources is like running a 6-cylinder engine on 2 spark plugs."
)

doc.add_heading('6.3 Storyline 3: Leadership (Sources 5 & 6 — Structural)', level=2)

add_para("Objective: Move from sponsorship-in-committee to visible, active modeling.", bold=True)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
headers = ["From (Cultural Debt)", "To (Cultural Infrastructure)", "Evidence"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

lead_data = [
    ["Incentives reward the old way",
     "Adoption tied to promotion criteria",
     "Financial services firm: responsible AI use in promotion criteria — adoption became a career advantage"],
    ["AI bolted on outside core workflows",
     "AI embedded into tools people already use",
     "Healthcare system: AI alerts in EMR workflow requiring review before chart close — adoption became automatic"],
    ["Sponsorship lives in a steering committee",
     "Executives use tools publicly, share failures",
     "MIT CISR: executive engagement is the strongest predictor of digital transformation success"],
    ["\"Bureaucratic antibodies\" — narrow scope, declare premature",
     "Structural redesign — remove friction, redesign around AI",
     "Ask: \"Would this step exist if we built the process from scratch?\" If not, it's cultural debt encoded in process."],
]
for r, row_data in enumerate(lead_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

add_para("", space_after=4)
add_quote(
    "If you're asking 2026 technology to operate inside a 2019 culture, "
    "the culture will win every time. — Shelly Palmer"
)

# ============================================================
# 7. PATH FORWARD
# ============================================================
doc.add_heading('7. Path Forward', level=1)

doc.add_heading('7.1 Three Actions This Quarter', level=2)

add_para(
    "Cultural debt cannot be eliminated in a quarter. But it can be named, measured, and "
    "systematically reduced. The following three actions — one per storyline — are designed "
    "to be low-cost, high-signal, and executable within 30 days."
)

doc.add_heading('Action 1: Audit the Cultural Balance Sheet (Culture — Sources 1 & 2)', level=3)

add_bullet("Run a cultural readiness assessment before the next AI investment.")
add_bullet("Survey team sentiment on AI: fear, trust, skill confidence, experimentation safety.")
add_bullet("Map the last 3 failed AI initiatives. Identify which of the five cultural patterns were present.")
add_bullet("Score the organization against the Cultural Debt Assessment Framework (see Appendix A).")
add_bullet("Owner: Executive Sponsor + HR. Timeline: 30 days.", bold_prefix="Owner: ")

doc.add_heading('Action 2: Realign the Scorecard (Change Management — Sources 3 & 4)', level=3)

add_bullet("Audit every AI program KPI. Flag any that measure deployment without measuring adoption.")
add_bullet("Replace deployment metrics with adoption and outcome metrics: time-to-value, weekly active users, decisions improved.")
add_bullet("Launch a cross-functional AI council with representation from Business, Technology, HR, and Compliance.")
add_bullet("Create peer support structures — communities of practice, shared prompt libraries, internal case studies.")
add_bullet("Owner: Business Unit Leads + Transformation Office. Timeline: 30 days.", bold_prefix="Owner: ")

doc.add_heading('Action 3: Make Leadership Visible (Leadership — Sources 5 & 6)', level=3)

add_bullet("One executive shares an AI experiment — success or failure — at the next all-hands. Publicly.")
add_bullet("Incorporate responsible AI use into the next performance review cycle as a leadership expectation.")
add_bullet("Identify one core workflow per business unit where AI can be embedded (not bolted on) and commit to integration.")
add_bullet("Owner: C-Suite + HR. Timeline: Next all-hands + next performance cycle.", bold_prefix="Owner: ")

doc.add_heading('7.2 Measuring Progress', level=2)

add_para(
    "Cultural debt requires its own metrics. Three proxies provide early signal:"
)

add_bullet(
    "Time from pilot to production adoption. If pilots succeed but production adoption "
    "stalls, the gap is cultural, not technical.",
    bold_prefix="Time-to-adoption gap: "
)
add_bullet(
    "Count the process steps between 'idea' and 'experiment.' Each step no one can explain "
    "is cultural debt encoded in process.",
    bold_prefix="Process step count: "
)
add_bullet(
    "Compare 'AI tools deployed' against 'AI tools used weekly.' The gap between these "
    "numbers is the cultural debt balance sheet.",
    bold_prefix="Deployed vs. used: "
)

doc.add_heading('7.3 The Governance Sequence', level=2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
headers = ["Phase", "Focus", "Timeline"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

gov_data = [
    ["Phase 1: Assess", "Cultural readiness assessment, baseline metrics, stakeholder alignment", "Q2 2026"],
    ["Phase 2: Intervene", "Scorecard realignment, AI councils, leadership visibility program", "Q3 2026"],
    ["Phase 3: Embed", "Cultural metrics in governance, adoption criteria in promotion cycles, process redesign", "Q4 2026"],
]
for r, row_data in enumerate(gov_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

add_para("", space_after=8)

add_para(
    "The path forward is not to slow down AI investment. It is to match every dollar of technology "
    "investment with a commensurate investment in the human system required to make it deliver value. "
    "AI-First defines what we build. People, Culture, and Mindset defines whether it gets adopted. "
    "An organization needs both.",
    bold=True
)

# ============================================================
# 8. REFERENCES
# ============================================================
doc.add_heading('8. References', level=1)

references = [
    "McKinsey Global Institute. (2023). The State of AI in 2023: Generative AI's Breakout Year. McKinsey & Company.",
    "Gartner. (2024). AI in the Enterprise Survey: Adoption, Challenges, and Outcomes. Gartner Research.",
    "BCG & MIT Sloan Management Review. (2024). Achieving Individual — and Organizational — Value with AI. Boston Consulting Group.",
    "Rand Corporation. (2024). Research Brief: Why AI Projects Fail. RAND Corporation.",
    "Deloitte. (2025). State of AI in the Enterprise, 7th Edition. Deloitte AI Institute.",
    "IDC. (2025). Worldwide Artificial Intelligence Spending Guide. International Data Corporation.",
    "De Cremer, D., Kasparov, G. et al. (2025). \"Why 95% of AI Initiatives Fail.\" Harvard Business Review.",
    "Grenny, J., Patterson, K. et al. (2008). \"How to 10x Your Influence.\" MIT Sloan Management Review.",
    "Crucial Learning. (2026). Influencer: The New Science of Leading Change, 3rd Edition.",
    "Palmer, S. (2026). \"AI Won't Fail Because of Technology.\" ShellyPalmer.com.",
    "MIT Center for Information Systems Research (CISR). (2024). Digital Transformation Research Series.",
    "Google re:Work. (2015). Project Aristotle: What Makes a Team Effective. Google.",
]

for ref in references:
    p = doc.add_paragraph(style='List Number')
    p.clear()
    run = p.add_run(ref)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

# ============================================================
# APPENDIX A: CULTURAL DEBT ASSESSMENT FRAMEWORK
# ============================================================
doc.add_page_break()
doc.add_heading('Appendix A: Cultural Debt Assessment Framework', level=1)

add_para(
    "Rate your organization 1–5 on each dimension. A total score below 15 signals significant "
    "cultural debt that must be addressed before scaling AI investment.",
    italic=True
)

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
headers = ["Dimension", "Storyline", "1 (High Debt)", "5 (Low Debt)"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

assess_data = [
    ["Experimentation Safety", "Culture", "Failures are punished or hidden", "Failures are shared and learned from openly"],
    ["Skill Building", "Culture", "Training is annual, classroom-based", "Learning is continuous, embedded in work"],
    ["Peer Support", "Change Mgmt", "Everyone figures it out alone", "AI councils and shared practices exist"],
    ["Adoption Metrics", "Change Mgmt", "Success = model deployed", "Success = business outcome improved"],
    ["Incentive Alignment", "Leadership", "Promotions reward pre-AI behaviors", "AI adoption is in promotion criteria"],
    ["Leadership Visibility", "Leadership", "Leaders talk about AI but don't use it", "Leaders demo their own AI usage regularly"],
]
for r, row_data in enumerate(assess_data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

add_para("", space_after=8)
add_para("Scoring:", bold=True)
add_bullet("25–30: Cultural infrastructure is strong — invest in technology with confidence.")
add_bullet("18–24: Moderate cultural debt — address gaps before scaling AI investment.")
add_bullet("12–17: Significant cultural debt — pause technology investment, invest in people first.")
add_bullet("6–11: Critical cultural debt — organizational redesign needed before AI can succeed.")

# ============================================================
# SAVE
# ============================================================
output_path = "/Users/GUNDLLX/learn-claude/Skill Test/presentation/why-ai-fails/Cultural-Debt-of-AI-Transformation-WhitePaper.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
