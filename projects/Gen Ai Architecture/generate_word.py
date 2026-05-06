"""
Abbott GenAI Blueprint — Word Document Generator
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = "/Users/GUNDLLX/learn-claude/projects/Gen Ai Architecture/Abbott-GenAI-Blueprint-Capability-Model-2026-05-06.docx"

NAVY_RGB = RGBColor(0x00, 0x00, 0x50)
MED_BLUE_RGB = RGBColor(0x33, 0x33, 0xAA)
LAVENDER_RGB = RGBColor(0xDD, 0xDD, 0xF8)
NEAR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)

DOMAIN_DATA = [
    {
        "name": "Domain 1: Content & Creative Studio",
        "maturity": "Scaling",
        "div_count": 7,
        "definition": "Generation, drafting, and transformation of written, visual, and multimedia content for internal and external audiences — spanning marketing copy, medical communications, training materials, and brand assets.",
        "subs": ["Creative Content Draft", "Writing Drafts & Data Entry", "Personalize Content", "Translation"],
        "tools": "Writer Newsroom (GMEA), Adobe Firefly/Express (RMDx), AutogenAI (AN)",
        "divisions": "ADC, Corp, EPD, GMEA, RMDx, CHR, MD",
        "examples": [
            "Personalized HCP marketing content (ADC)",
            "Medical affairs document generation (Corp)",
            "Field rep communication drafts (EPD)",
            "Multi-language regulatory translations (RMDx)",
        ],
    },
    {
        "name": "Domain 2: Insights & Analytics Acceleration",
        "maturity": "Producing",
        "div_count": 6,
        "definition": "AI-powered search, synthesis, and generation of insights from structured and unstructured data — including literature review, market intelligence, clinical data, and business reporting.",
        "subs": ["Search, Screen & Insights", "Reporting & Observability", "Forecast"],
        "tools": "Marketing Insights Accelerator (GMEA), RADIA Research (AN), Databricks (RMDx)",
        "divisions": "GMEA, AN, RMDx, BTS, ADC, Corp",
        "examples": [
            "Marketing Insights Accelerator (GMEA)",
            "Clinical trial literature screening (AN)",
            "Sales performance analytics (EPD, ADC)",
            "Adverse event signal detection (RMDx)",
        ],
    },
    {
        "name": "Domain 3: Productivity & Automation",
        "maturity": "Scaling",
        "div_count": 6,
        "definition": "Enterprise-wide productivity enhancement through AI-integrated tools — including meeting summarization, document management, workflow automation, and process intelligence embedded in everyday work tools.",
        "subs": ["Enterprise Productivity Tools", "Writing Drafts & Data Entry"],
        "tools": "Microsoft M365 Copilot (BTS), Cursor/Windsurf/Atlassian AI (BTS), Five9 GSD (BTS)",
        "divisions": "BTS, Corp, CHR, AN, MD, EPD",
        "examples": [
            "M365 Copilot enterprise rollout (BTS)",
            "HR process automation (CHR)",
            "Procurement document analysis (Corp)",
            "Audit workflow automation (Corp)",
        ],
    },
    {
        "name": "Domain 4: Customer & Field Engagement",
        "maturity": "Producing",
        "div_count": 5,
        "definition": "AI-powered tools that augment customer-facing roles — field sales, customer service, patient engagement, and HCP interactions — with intelligent assistants, next-best-action, and personalized recommendations.",
        "subs": ["Assistants", "Personalize Content", "Search & CRM Insights"],
        "tools": "SmartRep Call Planning (EPD), Neuron7 (CoreDx), Five9 GSD (BTS)",
        "divisions": "EPD, ADC, CoreDx, BTS, GMEA",
        "examples": [
            "SmartRep Call Planning for field reps (EPD)",
            "Neuron7 field service AI (CoreDx)",
            "Five9 AI-assisted customer service (BTS)",
            "HCP virtual assistant (ADC)",
        ],
    },
    {
        "name": "Domain 5: Engineering & Development Acceleration",
        "maturity": "Scaling",
        "div_count": 3,
        "definition": "AI tools that augment software engineering, quality assurance, and product development workflows — including code generation, testing, architecture assistance, and DevSecOps integration.",
        "subs": ["Code Creation with Platform", "Code Creation IDE & Test Authoring"],
        "tools": "Cursor/Windsurf/Atlassian AI (BTS), Azure AI Sandbox (MD)",
        "divisions": "BTS, MD, RMDx",
        "examples": [
            "Cursor/Windsurf AI-assisted coding (BTS)",
            "Atlassian AI for project management (BTS)",
            "Azure AI Sandbox for rapid prototyping (MD)",
            "Automated test generation (BTS)",
        ],
    },
    {
        "name": "Domain 6: Learning, Compliance & Governance",
        "maturity": "Exploring",
        "div_count": 6,
        "definition": "AI applications that support employee learning and development, regulatory compliance, risk management, legal review, and organizational knowledge management.",
        "subs": ["Learning, Training & Coaching", "Image Identification / Analysis"],
        "tools": "Pilots in CHR, Corp, RMDx, Cyber (multiple vendor-specific tools in pilot/WIP)",
        "divisions": "CHR, Corp, RMDx, MD, AN, Cyber",
        "examples": [
            "Personalized learning paths (CHR)",
            "Regulatory document review (Corp, RMDx)",
            "Legal contract analysis (Corp)",
            "Manufacturing defect image analysis (MD, AN)",
        ],
    },
]


def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def style_heading(para, text, level=1):
    run = para.add_run(text)
    run.font.color.rgb = NAVY_RGB
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)


def add_styled_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    n_rows = len(rows) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_widths:
        for i, w in enumerate(col_widths):
            table.columns[i].width = Inches(w)

    # Header row
    for ci, h in enumerate(headers):
        cell = table.cell(0, ci)
        cell.text = h
        set_cell_bg(cell, "000050")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(10)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            cell.text = str(val)
            if ri % 2 == 0:
                set_cell_bg(cell, "F0F0FF")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# ── Cover Page ────────────────────────────────────────────────
cover = doc.add_paragraph()
cover_run = cover.add_run("Abbott GenAI Blueprint & Capability Model")
cover_run.font.size = Pt(28)
cover_run.font.bold = True
cover_run.font.color.rgb = NAVY_RGB
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
sub = doc.add_paragraph()
sub_run = sub.add_run("Enterprise AI Landscape — 220 Use Cases Across 15 Divisions")
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = MED_BLUE_RGB
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
meta_table = doc.add_table(rows=4, cols=2)
meta_table.style = 'Table Grid'
meta_data = [
    ("Document Version", "1.0"),
    ("Date", "2026-05-06"),
    ("Owner", "Abbott BTS-DTS | GenAI Center of Excellence"),
    ("Classification", "Proprietary and Confidential — Do Not Distribute"),
]
for ri, (k, v) in enumerate(meta_data):
    meta_table.cell(ri, 0).text = k
    meta_table.cell(ri, 1).text = v
    for cell in [meta_table.cell(ri, 0), meta_table.cell(ri, 1)]:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
    set_cell_bg(meta_table.cell(ri, 0), "000050")
    for p in meta_table.cell(ri, 0).paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.bold = True

doc.add_page_break()

# ── 1. Executive Summary ─────────────────────────────────────
h = doc.add_paragraph()
style_heading(h, "1. Executive Summary", level=1)
doc.add_paragraph(
    "Abbott is executing a broad-based GenAI transformation across 12 business divisions, "
    "encompassing 220 catalogued use cases that span from clinical research to commercial marketing. "
    "As of 1Q26, 40 use cases are in production, 43 are in active development (WIP), and 29 are "
    "pending ESC approval. The portfolio reflects a deliberate Buy-first strategy (133 Buy vs. 86 Build), "
    "anchored on enterprise platforms such as Azure AI Foundry, Microsoft M365 Copilot, Salesforce "
    "Einstein, and Adobe Experience Cloud."
)
doc.add_paragraph(
    "The capability landscape has been consolidated from 15 raw capability categories into six "
    "strategic domains, providing a coherent blueprint for investment, governance, and platform "
    "standardization. The largest divisions by use case volume — BTS (32), ADC (31), and Corporate "
    "(27) — are driving enterprise-wide tooling, while commercial divisions like EPD and GMEA are "
    "leading production deployment in field engagement."
)

doc.add_page_break()

# ── 2. GenAI Landscape Overview ──────────────────────────────
h = doc.add_paragraph()
style_heading(h, "2. GenAI Landscape Overview", level=1)

h2 = doc.add_paragraph()
style_heading(h2, "2.1 Portfolio Statistics", level=2)
add_styled_table(doc,
    ["Metric", "Value", "Notes"],
    [
        ["Total use cases", "220", "Across all active, WIP, cancelled, and on-hold"],
        ["In production", "40 (18%)", "Generating live business value"],
        ["WIP / Active development", "43 (20%)", "Targeted for production in 2026"],
        ["Pending ESC approval", "29 (13%)", "Backlog to be cleared Q2 2026"],
        ["Cancelled / On Hold", "51 (23%)", "Deprioritized or deferred"],
        ["Buy vs. Build", "133 Buy / 86 Build", "61% Buy — platform-first strategy"],
        ["Divisions represented", "12", "BTS, ADC, Corp, AN, RMDx, MD, EPD, CHR, GMEA, CoreDx, Lingo, Cyber"],
        ["Top business function", "Marketing/Sales (22)", "Followed by Engineering (15) and HR (12)"],
    ],
    col_widths=[2.5, 2.0, 3.5]
)

doc.add_paragraph()
h2 = doc.add_paragraph()
style_heading(h2, "2.2 Domain Maturity Summary", level=2)
add_styled_table(doc,
    ["Domain", "Maturity", "In Production", "WIP", "Key Risk"],
    [
        ["Content & Creative Studio", "Scaling", "~8", "~12", "Platform fragmentation"],
        ["Insights & Analytics Acceleration", "Producing", "~14", "~9", "Data governance"],
        ["Productivity & Automation", "Scaling", "~10", "~11", "Change management"],
        ["Customer & Field Engagement", "Producing", "~12", "~6", "Compliance / regulatory"],
        ["Engineering & Development", "Scaling", "~6", "~8", "Talent / skill gaps"],
        ["Learning, Compliance & Governance", "Exploring", "~2", "~5", "Sensitivity / risk aversion"],
    ],
    col_widths=[3.0, 1.5, 1.5, 1.0, 2.5]
)

doc.add_page_break()

# ── 3. Six Capability Domains ─────────────────────────────────
h = doc.add_paragraph()
style_heading(h, "3. Six Capability Domains", level=1)
doc.add_paragraph(
    "Abbott's 15 raw GenAI capability categories have been logically grouped into six strategic "
    "domains that align with business outcomes, platform investments, and organizational structure."
)

for d in DOMAIN_DATA:
    doc.add_paragraph()
    h2 = doc.add_paragraph()
    style_heading(h2, d["name"], level=2)

    maturity_p = doc.add_paragraph()
    maturity_p.add_run("Maturity: ").bold = True
    maturity_p.add_run(d["maturity"])

    doc.add_paragraph(d["definition"])

    h3 = doc.add_paragraph()
    style_heading(h3, "Sub-capabilities & Details", level=3)

    add_styled_table(doc,
        ["Sub-Capability", "Status", "Notes"],
        [[s, "Active", ""] for s in d["subs"]],
        col_widths=[3.0, 1.5, 3.5]
    )

    doc.add_paragraph()
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ("In-Production Tools", d["tools"]),
        ("Divisions Using", d["divisions"]),
        ("Example Use Cases", " | ".join(d["examples"][:3])),
    ]
    for ri, (k, v) in enumerate(info_data):
        info_table.cell(ri, 0).text = k
        info_table.cell(ri, 1).text = v
        set_cell_bg(info_table.cell(ri, 0), "DDDDF8")
        for cell in [info_table.cell(ri, 0), info_table.cell(ri, 1)]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        for p in info_table.cell(ri, 0).paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = NAVY_RGB

doc.add_page_break()

# ── 4. Enterprise Blueprint ───────────────────────────────────
h = doc.add_paragraph()
style_heading(h, "4. Enterprise Blueprint — 3-Tier Architecture", level=1)
doc.add_paragraph(
    "Abbott's GenAI architecture follows a three-tier model that separates business outcomes, "
    "reusable AI capabilities, and shared platform infrastructure. This model enables CoE governance, "
    "cross-division reuse, and consistent security and compliance controls."
)

tiers = [
    ("Tier 1 — Business Domains",
     "The six capability domains represent business-outcome-oriented groupings of GenAI use cases. "
     "Each domain maps to one or more Abbott business functions and is owned by a domain champion "
     "within the CoE. Domains: Content & Creative Studio | Insights & Analytics Acceleration | "
     "Productivity & Automation | Customer & Field Engagement | Engineering & Development | "
     "Learning, Compliance & Governance"),
    ("Tier 2 — GenAI Capabilities",
     "Reusable AI capability patterns that can be shared across divisions and use cases. "
     "Sub-capabilities include: Creative Content Draft, Search & Insights, Enterprise Productivity "
     "Tools, Assistants & Personalization, Code Creation, and Learning & Compliance tools. "
     "These are governed by the CoE and documented in the capability catalog."),
    ("Tier 3 — Shared Platform Foundation",
     "Enterprise-licensed AI platforms that provide the underlying model access, infrastructure, "
     "and tooling. Platforms: Azure AI Foundry, Anthropic Claude, Salesforce Einstein, Adobe "
     "Experience Cloud, Veeva Vault AI, Microsoft M365 Copilot. All platforms are CoE governed "
     "with centralized access, billing, and security controls."),
]
for tier_name, tier_desc in tiers:
    h2 = doc.add_paragraph()
    style_heading(h2, tier_name, level=2)
    doc.add_paragraph(tier_desc)

doc.add_page_break()

# ── 5. Platform Foundation ────────────────────────────────────
h = doc.add_paragraph()
style_heading(h, "5. Platform Foundation", level=1)
add_styled_table(doc,
    ["Platform", "Primary Use Cases", "Domains Served", "Deployment Model"],
    [
        ["Azure AI Foundry", "Model access, AI infrastructure, custom fine-tuning", "All domains", "Enterprise"],
        ["Anthropic Claude", "Long-context reasoning, document analysis, code review", "Insights, Engineering, Compliance", "Enterprise API"],
        ["Salesforce Einstein", "CRM-native AI, sales copilot, next-best-action", "Customer & Field Engagement", "SaaS embedded"],
        ["Adobe Experience Cloud", "Content generation, creative tools, personalization", "Content & Creative Studio", "SaaS"],
        ["Veeva Vault AI", "Regulated content mgmt, MLR review, clinical docs", "Compliance & Governance", "SaaS"],
        ["Microsoft M365 Copilot", "Enterprise productivity, Teams, Office suite AI", "Productivity & Automation", "M365 tenant"],
    ],
    col_widths=[2.2, 3.5, 2.5, 1.8]
)

doc.add_page_break()

# ── 6. Division Breakdown ─────────────────────────────────────
h = doc.add_paragraph()
style_heading(h, "6. Division Breakdown", level=1)
add_styled_table(doc,
    ["Division", "Use Cases", "Top Domain", "In Production", "Notes"],
    [
        ["BTS", "32", "Productivity & Automation", "8", "M365 Copilot, Cursor, Five9"],
        ["ADC", "31", "Content & Creative", "6", "HCP marketing, personalization"],
        ["Corp", "27", "Insights & Analytics", "5", "Procurement, legal, audit"],
        ["AN", "22", "Insights & Analytics", "5", "RADIA Research, AutogenAI"],
        ["RMDx", "22", "Content & Creative", "4", "Databricks, Adobe Firefly"],
        ["MD", "22", "Engineering & Dev", "3", "Azure AI Sandbox"],
        ["EPD", "19", "Customer & Field", "4", "SmartRep Call Planning"],
        ["CHR", "14", "Learning & Governance", "2", "HR process automation"],
        ["GMEA", "9", "Customer & Field", "3", "Writer Newsroom, Mktg Insights"],
        ["CoreDx", "9", "Customer & Field", "3", "Neuron7"],
        ["Lingo", "5", "Productivity", "1", "Early stage"],
        ["Cyber", "5", "Compliance & Governance", "1", "Early stage"],
    ],
    col_widths=[1.2, 1.2, 2.5, 1.5, 3.0]
)

doc.add_page_break()

# ── 7. Recommendations & Next Steps ──────────────────────────
h = doc.add_paragraph()
style_heading(h, "7. Recommendations & Next Steps", level=1)
recommendations = [
    ("Consolidate Content Platforms",
     "Rationalize Writer, Adobe Firefly, and AutogenAI under a unified Content & Creative Studio "
     "platform contract with centralized governance. Target: single content AI platform by Q3 2026."),
    ("Accelerate ESC Approvals",
     "29 use cases are pending ESC — establish a 30-day SLA for review cadence to unblock deferred "
     "value. Assign dedicated ESC coordinator to track and route approvals."),
    ("Expand SmartRep Model to ADC & GMEA",
     "Template EPD's SmartRep success for replication across commercial divisions. Define data "
     "requirements, Salesforce Einstein integration playbook, and training program."),
    ("Launch Domain Champions Program",
     "Assign a CoE domain champion for each of the 6 capability domains to drive reuse, "
     "cross-division knowledge transfer, and quarterly domain reviews."),
    ("Formalize Build vs. Buy Criteria",
     "Define clear guardrails for when to Build vs. Buy to prevent shadow AI and unapproved "
     "custom models. Publish criteria to all divisions with CoE intake process."),
    ("Investment Plan for Learning & Governance Domain",
     "Develop a dedicated pilot use case in CHR to advance the Exploring → Scaling maturity level. "
     "Engage Cyber and Corp Legal/OEC as co-sponsors."),
]
for i, (title, desc) in enumerate(recommendations, 1):
    rh = doc.add_paragraph()
    style_heading(rh, f"{i}. {title}", level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ── Appendix: Capability Catalog ─────────────────────────────
h = doc.add_paragraph()
style_heading(h, "Appendix: GenAI Capability Catalog Summary", level=1)
add_styled_table(doc,
    ["Capability Category", "Count", "Domain", "Notes"],
    [
        ["Insights & Document Generation", "21", "Insights & Analytics", "Largest category"],
        ["Content / Document Creation", "19", "Content & Creative Studio", ""],
        ["Enterprise Productivity Tool", "19", "Productivity & Automation", ""],
        ["Virtual Assistants", "16", "Customer & Field Engagement", ""],
        ["Reporting", "14", "Insights & Analytics", ""],
        ["Search & Generate Insights/Reports", "14", "Insights & Analytics", ""],
        ["Creative Content Generation", "13", "Content & Creative Studio", ""],
        ["Process Automation", "10", "Productivity & Automation", ""],
        ["Translation", "6", "Content & Creative Studio", ""],
        ["Code Creation", "6", "Engineering & Development", ""],
        ["Interaction Chatbots", "5", "Customer & Field Engagement", ""],
        ["Image Identification/Analysis", "4", "Learning & Governance", ""],
        ["Coding Accelerators", "4", "Engineering & Development", ""],
        ["Learning, Training & Coaching", "3", "Learning & Governance", ""],
        ["Forecast", "1", "Insights & Analytics", "Nascent"],
    ],
    col_widths=[3.0, 1.0, 2.8, 2.5]
)

doc.save(OUT)
print(f"  Saved {OUT}")
