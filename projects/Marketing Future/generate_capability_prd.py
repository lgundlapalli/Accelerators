from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

NAVY = RGBColor(0x00, 0x00, 0x50)
MED_BLUE = RGBColor(0x33, 0x33, 0xAA)
LAVENDER_BG = RGBColor(0xDD, 0xDD, 0xF8)
NEAR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Page margins ────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# ── Style helpers ────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(18)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = MED_BLUE
        run.font.size = Pt(14)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(12)
    return p

def body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.color.rgb = NEAR_BLACK
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, 'DDDDFF')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = NAVY
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = NEAR_BLACK
    # Column widths
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t

def callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {text}  ")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    # Navy background via XML shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '000050')
    pPr.append(shd)
    return p

# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ABBOTT EPD MARKETING")
run.font.size = Pt(13)
run.font.color.rgb = LIGHT_GRAY
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EPD Marketer of the Future")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Capability PRD: Gap Analysis & Future State Roadmap")
run.font.size = Pt(16)
run.font.color.rgb = MED_BLUE

doc.add_paragraph()

meta = [
    ("Document Version", "1.0 — Investigation for Validation"),
    ("Date", "2026-05-05"),
    ("Status", "Draft — Pending Current State Validation"),
    ("Owner", "EPD Architect"),
    ("Source Documents", "The Marketer of the Future EPD 15.04.26.pdf | pharma-genai-tool-map[97].pdf"),
]
t = doc.add_table(rows=len(meta), cols=2)
t.style = 'Table Grid'
for i, (k, v) in enumerate(meta):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v
    set_cell_bg(t.rows[i].cells[0], 'DDDDFF')
    for p2 in t.rows[i].cells[0].paragraphs:
        for run in p2.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = NAVY
    for p2 in t.rows[i].cells[1].paragraphs:
        for run in p2.runs:
            run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════
h1(doc, "1. Executive Summary")
body(doc, "Abbott EPD Marketing is transforming from humans using AI tools to AI-empowered strategic orchestrators. This document defines 20 business capabilities across five purpose-led AI agents, maps current state (pending validation), identifies tool options, and outlines a validation workplan before architectural decisions are made.")
doc.add_paragraph()
body(doc, "Five systemic gaps drive this transformation:", bold=True)
for gap, desc in [
    ("Fragmentation", "Data, tools, and teams operate in silos with disconnected insight signals."),
    ("Reactive Mode", "Planning is static and annual; unable to pivot quickly to market."),
    ("Bottlenecks", "Manual production and review cycles create severe delays in output."),
    ("Mass Reach", "One-size-fits-all content due to inability to scale variations."),
    ("Compliance Gating", "MLR viewed as a gatekeeper rather than an integrated enabler."),
]:
    bullet(doc, f"{gap} — {desc}")

doc.add_paragraph()
callout(doc, "ALICE deployment proof point: brief quality from 6.2 to 8.9/10 | 95% faster execution | 750 days reclaimed | 3 weeks faster to market")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 2. FIVE AI AGENTS
# ═══════════════════════════════════════════════════════════
h1(doc, "2. Five Purpose-Led AI Agents")
body(doc, "The following five agents form an interconnected intelligence ecosystem covering the full EPD marketing lifecycle. Each agent targets specific systemic gaps.")
doc.add_paragraph()

add_table(doc,
    headers=["Agent", "Full Name", "Core Job", "Gaps Addressed"],
    rows=[
        ["ALICE", "AI-Led Intelligent Creative Engine", "Brief optimization, concept scoring, persona-targeted creative at scale", "Bottlenecks, Mass Reach"],
        ["BRANDS", "Brand Strategy Planning System", "Global brand plan, regional localization, market archetype navigation", "Fragmentation, Reactive Mode"],
        ["CRAFTS", "Content Lifecycle Management Agent", "Modular assets, MLR acceleration, pre-launch content testing", "Bottlenecks, Compliance Gating"],
        ["SPARK", "Intelligent Customer Engagement Agent", "NBA, HCP segmentation, omnichannel personalization, patient engagement", "Mass Reach, Fragmentation"],
        ["ACCESS", "AI-Powered Market Access Agent", "Competitive intelligence, HTA dossiers, patient support automation", "Reactive Mode, Compliance Gating"],
    ],
    col_widths=[1.0, 2.2, 4.0, 2.5]
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 3. CAPABILITY GAP ANALYSIS
# ═══════════════════════════════════════════════════════════
h1(doc, "3. Capability Gap Analysis")
body(doc, "The table below maps all 20 capabilities across the five agents. Current state tools are marked as 'Investigation for Validation' — a tool inventory discovery workplan must be completed before confirmed gap classification. See Section 5 for the validation workplan.")
doc.add_paragraph()

agents = [
    {
        "name": "3.1  ALICE — Creative Excellence",
        "subtitle": "AI-Led Intelligent Creative Engine: brief optimization, concept scoring, persona-targeted creative at scale",
        "pain_points": [
            "Manual, inconsistent creative briefs (agency-dependent, high cost, long lead times)",
            "3–5 day brief cycle with 4–6 revision cycles",
            "Intuition-based judgments with limited localization",
            "Linear campaign rollouts, slow time to market",
        ],
        "caps": [
            ["1", "Creative Generation", "Generate brand-safe images, video, and modular assets at scale", "Investigation for Validation", "Adobe Firefly / GenStudio (Market Leader)\nMidjourney / Runway Gen-4 (Emerging)", "TBD"],
            ["2", "Brief Intelligence", "Auto-optimize creative briefs; score concepts against effectiveness benchmarks", "Investigation for Validation", "Persado (Market Leader)\nJasper AI (Core Enablement)", "TBD"],
            ["3", "Copy & Claim Generation", "Generate pharma-compliant copy variants with brand voice guardrails", "Investigation for Validation", "Writer.com (Market Leader)\nJasper AI (Core Enablement)", "TBD"],
            ["4", "Creative Ops & Workflow", "Connect brief intake to asset production with automated routing and review", "Investigation for Validation", "Adobe Workfront (Core Enablement)", "TBD"],
        ]
    },
    {
        "name": "3.2  BRANDS — Brand Strategy",
        "subtitle": "Brand Strategy Planning System: builds global brand plan, localizes for markets, deploys at scale",
        "pain_points": [
            "Static insights from fragmented data sources",
            "Resource-intensive 2–3 month brand workshops",
            "Annual brand plans with limited adaptability to market changes",
            "Intuition-based judgments rather than data-driven decisions",
        ],
        "caps": [
            ["5", "Market Intelligence", "Synthesize research, competitive intel, HCP/patient sentiment into brand strategy", "Investigation for Validation", "Deepsights (Market Leader)\nIQVIA HCP Data + AI Insights (Core Enablement)", "TBD"],
            ["6", "Social Listening", "Real-time brand perception monitoring across HCP and patient communities", "Investigation for Validation", "Brandwatch / Synthesio (Market Leader)", "TBD"],
            ["7", "Commercial Strategy Modeling", "Scenario modeling, portfolio planning, and market archetype analysis", "Investigation for Validation", "McKinsey Lilli / ZS ZAIDYN (Emerging)\nSalesforce Einstein LS Cloud (Core Enablement)", "TBD"],
            ["8", "Synthetic Research / Digital Twins", "Generate synthetic HCP & patient personas for rapid message testing without traditional research", "Investigation for Validation", "Synthetic Audiences / Digital Twins (Emerging)", "TBD"],
        ]
    },
    {
        "name": "3.3  CRAFTS — Content Excellence",
        "subtitle": "Content Lifecycle Management Agent: conceives strategy, renders modular assets, manages MLR, tests before launch",
        "pain_points": [
            "Slow MLR reviews due to manual coordination bottlenecks",
            "Linear campaign rollouts — asset-based, not customer journey-based",
            "One-size-fits-all campaigns, no content variation at scale",
            "Compliance viewed as a gatekeeper, not an integrated enabler",
        ],
        "caps": [
            ["9", "MLR Review & Compliance", "AI-assisted pre-review, claim checking, citation linking, risk flagging", "Investigation for Validation", "Veeva Vault PromoMats + AI Agents (MLR-Critical)\nPapercurve (Emerging)", "TBD"],
            ["10", "Compliance Detection", "Auto-detect non-compliant claims and flag risk in global content audits", "Investigation for Validation", "Indegene AI Compliance Engine (MLR-Critical)", "TBD"],
            ["11", "Modular Content Creation", "Generate claim-safe modular content blocks for multi-channel asset rendering", "Investigation for Validation", "Writer.com (Market Leader)\nAdobe Experience Manager + Firefly (Market Leader)", "TBD"],
            ["12", "Content Testing & Optimization", "A/B and multivariate testing for HCP-facing digital content before mass go-live", "Investigation for Validation", "Optimizely / VWO AI (Core Enablement)", "TBD"],
        ]
    },
    {
        "name": "3.4  SPARK — Customer Engagement",
        "subtitle": "Intelligent Customer Engagement Agent: segments audiences, prioritizes targets, activates channels, routes personalized journeys",
        "pain_points": [
            "Mass reach — one-size-fits-all content, unable to scale personalization",
            "Broad segments, siloed data, no real-time signal processing",
            "Reactive planning mode — static, not responsive to market dynamics",
            "Disconnected tools preventing unified HCP and patient journeys",
        ],
        "caps": [
            ["13", "Next-Best-Action (NBA)", "AI-driven field suggestions for rep interactions; HCP pre-call intelligence and rep coaching", "Investigation for Validation", "Aktana / PharmaForceIQ (Market Leader, merged Jan 2026)\nVerix Sage (Emerging)", "TBD"],
            ["14", "CRM & HCP Journey Orchestration", "Einstein AI-powered HCP segmentation, behavioral analytics, and omnichannel campaign delivery", "Investigation for Validation", "Veeva Vault CRM AI Agents (Market Leader)\nSalesforce Life Sciences Cloud (Core Enablement)", "TBD"],
            ["15", "Patient Engagement", "Conversational AI handling HCP/patient queries on dosing, side effects, and access", "Investigation for Validation", "Conversa Health (Emerging)\nAdobe Sensei / Marketing Cloud (Core Enablement)", "TBD"],
            ["16", "Omnichannel Personalization", "Real-time content personalization and send-time optimization across email, web, and field channels", "Investigation for Validation", "Adobe Sensei / Marketing Cloud (Core Enablement)", "TBD"],
        ]
    },
    {
        "name": "3.5  ACCESS — Market Access",
        "subtitle": "AI-Powered Market Access Agent: analyzes competition, builds value dossiers, simulates economic models, supports patient programs",
        "pain_points": [
            "Reactive payer response — post-campaign analysis with limited actionability",
            "Siloed data and quantitative-only measurement",
            "Manual HTA dossier and economic model preparation (months, not weeks)",
            "Compliance managed separately from market access strategy",
        ],
        "caps": [
            ["17", "Competitive Intelligence", "Monitor competitor launches, payer decisions, and access signals in real time", "Investigation for Validation", "IQVIA Market Edge AI (Market Leader)\nCrayon / Klue AI (Emerging)", "TBD"],
            ["18", "HTA & Value Dossier", "Build and simulate economic value dossiers, ICER models, HTA submission narratives", "Investigation for Validation", "Payer Sciences / ZS ZAIDYN Access (Emerging)\nAvalere AI / Inovalon (Core Enablement)", "TBD"],
            ["19", "Patient Support Programs", "Automate copay, prior authorization, and patient assistance program operations", "Investigation for Validation", "ConnectiveRx AI Platform (Market Leader)", "TBD"],
            ["20", "Regulatory & Access Compliance", "Scan content for access-related risk, OPDP compliance exposure, and global claim audit trails", "Investigation for Validation", "RegASK / Veripharm (MLR-Critical)", "TBD"],
        ]
    },
]

for agent in agents:
    h2(doc, agent["name"])
    body(doc, agent["subtitle"])
    doc.add_paragraph()
    body(doc, "Known Pain Points:", bold=True)
    for pt in agent["pain_points"]:
        bullet(doc, pt)
    doc.add_paragraph()
    add_table(doc,
        headers=["#", "Capability", "Definition", "Current State", "Future Tool Options", "Gap Status"],
        rows=agent["caps"],
        col_widths=[0.3, 1.5, 2.5, 1.5, 2.5, 1.0]
    )
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 4. TOOL CLASSIFICATION GUIDE
# ═══════════════════════════════════════════════════════════
h1(doc, "4. Tool Classification Guide")
body(doc, "Tools are classified by adoption readiness to guide sequencing of evaluation and procurement.")
doc.add_paragraph()
add_table(doc,
    headers=["Classification", "Definition", "Examples from This PRD"],
    rows=[
        ["Market Leader", "Proven at scale in pharma; highest confidence for Abbott EPD adoption", "Persado, Writer.com, Adobe Firefly, Aktana/PharmaForceIQ, IQVIA Market Edge AI, ConnectiveRx"],
        ["Core Enablement", "Strong platforms with AI features; may already be licensed at Abbott", "Salesforce Einstein LS Cloud, Adobe Sensei/Marketing Cloud, Veeva Vault CRM, Optimizely"],
        ["Emerging", "High-potential 2026 tools; validate before committing; strong ROI signals", "McKinsey Lilli, ZS ZAIDYN, Conversa Health, Synthetic Audiences, Crayon/Klue AI"],
        ["MLR-Critical", "Regulatory workflow tools; require Medical/Regulatory sign-off before adoption", "Veeva Vault PromoMats + AI Agents, Indegene AI, Papercurve, RegASK / Veripharm"],
    ],
    col_widths=[1.5, 3.5, 4.5]
)
doc.add_paragraph()
callout(doc, "2026 Priority Trends: MLR Transformation | Synthetic Audiences Rising | Signal-to-Action Intelligence — 38% of MLR projected to be AI-driven by 2028")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 5. VALIDATION WORKPLAN
# ═══════════════════════════════════════════════════════════
h1(doc, "5. Validation Workplan")
body(doc, "Before any tool procurement decisions are made, the following discovery activities must be completed. The fastest path to the right technology decisions is knowing what Abbott EPD already owns.")
doc.add_paragraph()
add_table(doc,
    headers=["Priority", "Action", "Owner", "Timeline", "Why It Matters"],
    rows=[
        ["P1", "Inventory current MarTech tools across all EPD regions", "Marketing Ops", "Week 1–4", "Avoid buying what we already own"],
        ["P1", "Identify existing Veeva, Salesforce, and Adobe enterprise licenses", "IT / Procurement", "Week 1–4", "Major platforms may already be licensed — unlocks Core Enablement tools at no additional cost"],
        ["P1", "Map current MLR process and tooling across markets", "Medical / Regulatory", "Week 1–4", "MLR is the highest-impact gap to close; determines CRAFTS architecture"],
        ["P2", "Assess CRM maturity (Veeva vs Salesforce vs other)", "Commercial IT", "Week 5–6", "CRM platform determines SPARK agent architecture and data flow"],
        ["P2", "Review existing analytics and data platforms", "Data & Analytics", "Week 5–6", "Data foundation quality drives the effectiveness of all 5 agents"],
        ["P3", "Identify regional variation in tool usage across EPD markets", "Regional Marketing Leads", "Week 7–8", "EPD markets may differ significantly — a single global tool stack may not apply"],
    ],
    col_widths=[0.8, 3.0, 2.0, 1.2, 2.5]
)
doc.add_paragraph()
body(doc, "Validation Sequence:")
bullet(doc, "P1 actions (Weeks 1–4): Tool inventory, license audit, MLR process mapping")
bullet(doc, "P2 actions (Weeks 5–6): CRM assessment, data platform review")
bullet(doc, "P3 actions (Weeks 7–8): Regional variation analysis")
bullet(doc, "Output (Week 9–10): Updated gap analysis with confirmed current state, confirmed gap types, and prioritized tool recommendations")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 6. NEXT STEPS
# ═══════════════════════════════════════════════════════════
h1(doc, "6. Next Steps & Recommended Actions")
body(doc, "Three actions move this assessment from a capability map to a buildable architecture.")
doc.add_paragraph()
add_table(doc,
    headers=["#", "Action", "Owner", "Timeline", "Output"],
    rows=[
        ["1", "Complete the Validation Workplan (Section 5)", "Marketing Ops + IT/Procurement + Medical/Regulatory", "4 weeks", "Confirmed current state tool inventory; updated gap analysis"],
        ["2", "Run Current State → Future State Capability Mapping", "EPD Architect + Marketing Ops", "2 weeks post-validation", "Validated gap analysis with confirmed gap types and tool shortlist"],
        ["3", "Build Integration & Target Architecture", "EPD Architect", "4 weeks post-mapping", "Integration architecture + target architecture diagrams for all 5 agents"],
    ],
    col_widths=[0.4, 3.2, 2.4, 1.5, 2.0]
)
doc.add_paragraph()
callout(doc, "Validation unlocks mapping. Mapping unlocks architecture. Architecture unlocks procurement. Follow the sequence.")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 7. APPENDIX
# ═══════════════════════════════════════════════════════════
h1(doc, "Appendix A: Full Tool Reference")
body(doc, "Complete list of all 28 tools evaluated across the five agents.")
doc.add_paragraph()
add_table(doc,
    headers=["Tool", "Agent", "Category", "Classification", "Key Capability"],
    rows=[
        ["Adobe Firefly / GenStudio", "ALICE", "Creative Generation", "Market Leader", "Brand-safe image & video generation; modular asset creation in Adobe workflows"],
        ["Writer.com", "ALICE / CRAFTS", "Copy & Content", "Market Leader", "Enterprise GenAI with brand voice guardrails and claim-safe copy; pharma-tunable LLM"],
        ["Persado", "ALICE", "Brief Intelligence", "Market Leader", "Emotion AI for messaging optimization; generates compliant copy variants at scale"],
        ["Jasper AI", "ALICE", "Brief & Copy", "Core Enablement", "AI writing for brief creation, headline generation, campaign copy iteration"],
        ["Adobe Workfront", "ALICE", "Creative Ops", "Core Enablement", "AI-assisted project management; connects brief intake to asset production"],
        ["Midjourney / Runway Gen-4", "ALICE", "Creative Generation", "Emerging", "High-fidelity visual ideation, mood boards, motion assets for campaign concepting"],
        ["Deepsights", "BRANDS", "Market Intelligence", "Market Leader", "AI-powered synthesis of research, competitive intel, HCP/patient sentiment"],
        ["Brandwatch / Synthesio", "BRANDS", "Social Listening", "Market Leader", "Social listening + GenAI narrative analysis for real-time brand perception monitoring"],
        ["McKinsey Lilli / ZS ZAIDYN", "BRANDS", "Strategy Modeling", "Emerging", "Enterprise AI for commercial strategy scenario modeling and portfolio planning"],
        ["Salesforce Einstein LS Cloud", "BRANDS / SPARK", "CRM & Analytics", "Core Enablement", "Predictive analytics and AI-driven segmentation for HCP positioning and journeys"],
        ["Synthetic Audiences / Digital Twins", "BRANDS", "Research", "Emerging", "GenAI-generated synthetic HCP & patient personas for rapid message testing"],
        ["IQVIA HCP Data + AI Insights", "BRANDS", "Market Intelligence", "Core Enablement", "Real-world data enriched with AI for prescriber behavior and brand signals"],
        ["Veeva Vault PromoMats + AI Agents", "CRAFTS", "MLR", "MLR-Critical", "Industry-standard MLR platform with Quick Check Agent; up to 75% cycle time reduction"],
        ["Indegene AI Compliance Engine", "CRAFTS", "Compliance", "MLR-Critical", "Auto-detects non-compliant claims; links citations; flags risk in global audits"],
        ["Adobe Experience Manager + Firefly", "CRAFTS", "Content Creation", "Market Leader", "AI-assisted modular content system for multi-channel asset rendering at scale"],
        ["Papercurve", "CRAFTS", "MLR Pre-Review", "Emerging", "Auto-annotates claims against references for MLR pre-submission speed-ups"],
        ["Optimizely / VWO AI", "CRAFTS", "Content Testing", "Core Enablement", "GenAI-powered A/B and multivariate testing for HCP-facing digital content"],
        ["Aktana / PharmaForceIQ", "SPARK", "Next-Best-Action", "Market Leader", "Industry-leading NBA platform (merged Jan 2026); 36% NBRx lift; 100M+ field suggestions"],
        ["Veeva Vault CRM AI Agents", "SPARK", "CRM", "Market Leader", "Pre-call Agent, Voice Agent, Free Text Agent for HCP intelligence and rep coaching"],
        ["Salesforce Life Sciences Cloud", "SPARK", "CRM / Orchestration", "Core Enablement", "Einstein AI CRM with HCP journey orchestration and omnichannel delivery (Sep 2025)"],
        ["Verix Sage", "SPARK", "NBA / Engagement", "Emerging", "LLM-powered HCP engagement intelligence translating data signals to pre-call strategies"],
        ["Conversa Health", "SPARK", "Patient Engagement", "Emerging", "Conversational AI handling 60%+ of HCP/patient queries on dosing, side effects, access"],
        ["Adobe Sensei / Marketing Cloud", "SPARK", "Personalization", "Core Enablement", "Real-time content personalization and send-time optimization across email and web"],
        ["IQVIA Market Edge AI", "ACCESS", "Competitive Intel", "Market Leader", "Real-world evidence and market dynamics analysis; identifies payer landscape shifts"],
        ["Crayon / Klue AI", "ACCESS", "Competitive Intel", "Emerging", "GenAI-powered competitive intelligence monitoring competitor launches and payer decisions"],
        ["Payer Sciences / ZS ZAIDYN Access", "ACCESS", "HTA & Dossier", "Emerging", "AI tools for economic value dossiers, ICER models, HTA submission narratives"],
        ["Avalere AI (Inovalon)", "ACCESS", "HTA & Policy", "Core Enablement", "Policy intelligence and payer coverage modeling with GenAI scenario simulation"],
        ["ConnectiveRx AI Platform", "ACCESS", "Patient Support", "Market Leader", "GenAI-optimized patient services — copay, prior auth, patient assistance automation"],
        ["RegASK / Veripharm", "ACCESS", "Regulatory Compliance", "MLR-Critical", "Regulatory AI scanning content for access-related risk and OPDP compliance exposure"],
    ],
    col_widths=[2.0, 1.0, 1.5, 1.3, 3.7]
)

# ── Save ─────────────────────────────────────────────────────
output_path = "/Users/GUNDLLX/learn-claude/projects/Marketing Future/EPD-Marketer-of-the-Future-Capability-PRD-2026-05-05.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
